"""
Text Extraction Module
Extracts text from PDF, DOCX, and image files.
- PDF: pdfplumber for layout-aware extraction
- DOCX: python-docx for paragraphs and headings
- Image: OpenCV preprocessing + Tesseract OCR
"""
import io
import os
import sys
import io
import os
import sys
from docx import Document

# Configure Tesseract path for Windows
if sys.platform == "win32":
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for tp in tesseract_paths:
        if os.path.exists(tp):
            pytesseract.pytesseract.tesseract_cmd = tp
            break


def extract_from_pdf(file_bytes: bytes) -> str:
    raise ValueError("PDFs should now be processed by multimodal Gemini endpoint.")


def extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX preserving paragraphs and headings."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Preserve heading structure
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading", "").strip()
            text_parts.append(f"\n{'#' * int(level) if level.isdigit() else '#'} {text}\n")
        else:
            text_parts.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)
    
    return "\n".join(text_parts)


def extract_from_image(file_bytes: bytes) -> str:
    raise ValueError("Images should now be processed by multimodal Gemini endpoint.")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Route to appropriate text extractor based on file type.
    
    Args:
        file_bytes: Raw file content
        file_type: 'pdf', 'docx', or 'image'
    
    Returns:
        Extracted text string
    """
    extractors = {
        "pdf": extract_from_pdf,
        "docx": extract_from_docx,
        "image": extract_from_image,
    }
    
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"No extractor available for file type: {file_type}")
    
    text = extractor(file_bytes)
    
    if not text or not text.strip():
        raise ValueError(
            f"No text could be extracted from the {file_type} file. "
            "The file may be empty, corrupted, or contain only non-text elements."
        )
    
    return text
