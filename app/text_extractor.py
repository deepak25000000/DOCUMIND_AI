"""
Text Extraction Module
Extracts text from PDF, DOCX, and image files.
- PDF: pdfplumber for layout-aware extraction
- DOCX: python-docx for paragraphs and headings
- Image: OpenCV preprocessing + Tesseract OCR
"""
import io
import os
import sys
import numpy as np
from PIL import Image
import cv2
import pytesseract
import pdfplumber
from docx import Document

# Configure Tesseract path for Windows
if sys.platform == "win32":
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for tp in tesseract_paths:
        if os.path.exists(tp):
            pytesseract.pytesseract.tesseract_cmd = tp
            break


def extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF preserving layout structure."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            
            # Also try extracting from tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        row_text = " | ".join(
                            cell.strip() if cell else "" for cell in row
                        )
                        if row_text.strip(" |"):
                            text_parts.append(row_text)
    
    return "\n\n".join(text_parts)


def extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX preserving paragraphs and headings."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Preserve heading structure
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading", "").strip()
            text_parts.append(f"\n{'#' * int(level) if level.isdigit() else '#'} {text}\n")
        else:
            text_parts.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)
    
    return "\n".join(text_parts)


def preprocess_image(image: Image.Image) -> np.ndarray:
    """
    Preprocess image for better OCR accuracy.
    - Convert to grayscale
    - Noise reduction
    - Adaptive thresholding
    """
    # Convert PIL Image to numpy array
    img_array = np.array(image)
    
    # Convert to grayscale if needed
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    else:
        gray = img_array
    
    # Noise reduction using Gaussian blur  
    denoised = cv2.GaussianBlur(gray, (3, 3), 0)
    
    # Adaptive thresholding for better text contrast
    thresh = cv2.adaptiveThreshold(
        denoised, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 11, 2
    )
    
    return thresh


def extract_from_image(file_bytes: bytes) -> str:
    """Extract text from image using Tesseract OCR with preprocessing."""
    # Open image
    image = Image.open(io.BytesIO(file_bytes))
    
    # Convert to RGB if necessary (handle RGBA, palette, etc.)
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")
    
    # Preprocess for better OCR
    processed = preprocess_image(image)
    
    # Convert back to PIL Image for pytesseract
    processed_image = Image.fromarray(processed)
    
    # Run Tesseract OCR with optimized config
    custom_config = r'--oem 3 --psm 6'
    text = pytesseract.image_to_string(processed_image, config=custom_config)
    
    return text


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Route to appropriate text extractor based on file type.
    
    Args:
        file_bytes: Raw file content
        file_type: 'pdf', 'docx', or 'image'
    
    Returns:
        Extracted text string
    """
    extractors = {
        "pdf": extract_from_pdf,
        "docx": extract_from_docx,
        "image": extract_from_image,
    }
    
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"No extractor available for file type: {file_type}")
    
    text = extractor(file_bytes)
    
    if not text or not text.strip():
        raise ValueError(
            f"No text could be extracted from the {file_type} file. "
            "The file may be empty, corrupted, or contain only non-text elements."
        )
    
    return text
